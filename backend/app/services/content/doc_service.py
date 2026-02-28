import os
import uuid
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.services.ai.content_planner import plan_document


OUTPUT_DIR = "generated_files"
os.makedirs(OUTPUT_DIR, exist_ok=True)


def set_heading_style(paragraph, level: int = 1):
    """Apply heading style to paragraph"""
    paragraph.style = f'Heading {level}'
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    run.font.bold = True
    run.font.size = Pt(16 if level == 1 else 13)


def add_cover_page(doc: Document, title: str, subject: str, student_name: str, college_name: str):
    """Add a professional cover page"""

    # College name
    college_para = doc.add_paragraph()
    college_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    college_run = college_para.add_run(college_name or "EduAI Suite")
    college_run.font.size = Pt(18)
    college_run.font.bold = True
    college_run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    doc.add_paragraph()
    doc.add_paragraph()

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run(title)
    title_run.font.size = Pt(28)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

    doc.add_paragraph()

    # Subject
    if subject:
        sub_para = doc.add_paragraph()
        sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sub_run = sub_para.add_run(f"Subject: {subject}")
        sub_run.font.size = Pt(14)
        sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # Student name
    if student_name:
        name_para = doc.add_paragraph()
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_run = name_para.add_run(f"Submitted by: {student_name}")
        name_run.font.size = Pt(13)
        name_run.font.bold = True

    # Page break after cover
    doc.add_page_break()


def generate_document(
    title: str,
    topic: str,
    doc_type: str = "assignment",
    key_points: list = None,
    subject: str = None,
    student_name: str = None,
    college_name: str = None,
    include_references: bool = True
) -> dict:
    """
    Generate a complete Word document.

    1. LangChain plans content
    2. python-docx builds the file
    3. Saves to disk
    4. Returns file path
    """

    print(f"📄 Generating document: {title}")

    # Step 1 — Plan content with AI
    content = plan_document(
        title=title,
        topic=topic,
        doc_type=doc_type,
        key_points=key_points,
        subject=subject
    )

    # Step 2 — Create Word document
    doc = Document()

    # Page margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)

    # Step 3 — Cover page
    add_cover_page(doc, title, subject, student_name, college_name)

    # Step 4 — Abstract
    if content.get("abstract"):
        abs_heading = doc.add_heading("Abstract", level=1)
        abs_heading.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
        abs_para = doc.add_paragraph(content["abstract"])
        abs_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_paragraph()

    # Step 5 — Introduction
    if content.get("introduction"):
        intro_heading = doc.add_heading("1. Introduction", level=1)
        intro_heading.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
        intro_para = doc.add_paragraph(content["introduction"])
        intro_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_paragraph()

    # Step 6 — Main sections
    for i, section in enumerate(content.get("sections", []), start=2):
        heading = doc.add_heading(f"{i}. {section.get('heading', 'Section')}", level=1)
        heading.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
        content_para = doc.add_paragraph(section.get("content", ""))
        content_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_paragraph()

    # Step 7 — Conclusion
    if content.get("conclusion"):
        next_num = len(content.get("sections", [])) + 2
        conc_heading = doc.add_heading(f"{next_num}. Conclusion", level=1)
        conc_heading.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
        conc_para = doc.add_paragraph(content["conclusion"])
        conc_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        doc.add_paragraph()

    # Step 8 — References
    if include_references and content.get("references"):
        ref_heading = doc.add_heading("References", level=1)
        ref_heading.runs[0].font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
        for i, ref in enumerate(content["references"], 1):
            doc.add_paragraph(f"[{i}] {ref}")

    # Step 9 — Save file
    file_id = str(uuid.uuid4())
    file_name = f"{file_id}_{title[:30].replace(' ', '_')}.docx"
    file_path = os.path.join(OUTPUT_DIR, file_name)
    doc.save(file_path)

    # Count words
    word_count = sum(
        len(str(v).split())
        for v in content.values()
        if isinstance(v, str)
    )

    print(f"✅ Document saved: {file_path}")

    return {
        "file_path": file_path,
        "file_name": file_name,
        "word_count": word_count
    }