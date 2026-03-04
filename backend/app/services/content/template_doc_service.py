import os
import uuid
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from app.schemas.template_schema import CollegeTemplate
from app.services.ai.content_planner import plan_document

OUTPUT_DIR = "generated_files"
os.makedirs(OUTPUT_DIR, exist_ok=True)


def hex_to_rgb_tuple(hex_color: str) -> tuple:
    """Convert hex to (r, g, b) tuple"""
    hex_color = hex_color.lstrip('#')
    return tuple(int(hex_color[i:i+2], 16) for i in (0, 2, 4))


def _add_college_header(doc: Document, template: CollegeTemplate):
    """
    Add a professional college header to the document.
    Includes college name, divider line, department.
    """
    section = doc.sections[0]
    header = section.header

    # College name — centered, large, colored
    college_para = header.paragraphs[0]
    college_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r, g, b = hex_to_rgb_tuple(template.primary_color)
    run = college_para.add_run(template.college_name)
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.name = template.font_name
    run.font.color.rgb = RGBColor(r, g, b)

    # Department / footer_text line below college name
    if template.footer_text:
        dept_para = header.add_paragraph()
        dept_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dept_run = dept_para.add_run(template.footer_text)
        dept_run.font.size = Pt(11)
        dept_run.font.name = template.font_name
        r2, g2, b2 = hex_to_rgb_tuple(template.secondary_color)
        dept_run.font.color.rgb = RGBColor(r2, g2, b2)

    # Divider line
    divider_para = header.add_paragraph()
    divider_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    divider_para.paragraph_format.space_after = Pt(0)

    # Add bottom border to divider paragraph
    pPr = divider_para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    r3, g3, b3 = hex_to_rgb_tuple(template.primary_color)
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), template.primary_color.lstrip('#'))
    pBdr.append(bottom)
    pPr.append(pBdr)


def _add_college_footer(doc: Document, template: CollegeTemplate):
    """Add page number and college name to footer"""
    section = doc.sections[0]
    footer = section.footer

    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    r, g, b = hex_to_rgb_tuple(template.secondary_color)

    left_run = footer_para.add_run(f"{template.college_name}  |  ")
    left_run.font.size = Pt(9)
    left_run.font.name = template.font_name
    left_run.font.color.rgb = RGBColor(r, g, b)

    # Add auto page number field
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run_xml = footer_para.add_run()
    run_xml._r.append(fldChar1)
    run_xml._r.append(instrText)
    run_xml._r.append(fldChar2)
    run_xml.font.size = Pt(9)
    run_xml.font.color.rgb = RGBColor(r, g, b)


def _add_college_cover_page(
    doc: Document,
    title: str,
    subject: str,
    student_name: str,
    doc_type: str,
    template: CollegeTemplate
):
    """Add a branded college cover page"""
    r, g, b = hex_to_rgb_tuple(template.primary_color)
    r2, g2, b2 = hex_to_rgb_tuple(template.secondary_color)

    # ── College name (big, top center) ──────────────────
    college_para = doc.add_paragraph()
    college_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    college_para.paragraph_format.space_before = Pt(30)
    cr = college_para.add_run(template.college_name)
    cr.font.size = Pt(22)
    cr.font.bold = True
    cr.font.name = template.font_name
    cr.font.color.rgb = RGBColor(r, g, b)

    # ── Department ──────────────────
    if template.footer_text:
        dept_para = doc.add_paragraph()
        dept_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        dr = dept_para.add_run(template.footer_text)
        dr.font.size = Pt(14)
        dr.font.name = template.font_name
        dr.font.color.rgb = RGBColor(r2, g2, b2)

    # ── Colored divider ──────────────────
    div_para = doc.add_paragraph()
    div_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    div_para.paragraph_format.space_before = Pt(8)
    div_para.paragraph_format.space_after = Pt(8)
    pPr = div_para._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '12')
    bottom.set(qn('w:color'), template.primary_color.lstrip('#'))
    pBdr.append(bottom)
    pPr.append(pBdr)
    doc.add_paragraph()

    # ── Document type label ──────────────────
    type_para = doc.add_paragraph()
    type_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = type_para.add_run(doc_type.replace("_", " ").upper())
    tr.font.size = Pt(13)
    tr.font.name = template.font_name
    tr.font.color.rgb = RGBColor(r2, g2, b2)
    tr.font.bold = True

    doc.add_paragraph()
    doc.add_paragraph()

    # ── Main title ──────────────────
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titler = title_para.add_run(title)
    titler.font.size = Pt(26)
    titler.font.bold = True
    titler.font.name = template.font_name
    titler.font.color.rgb = RGBColor(r, g, b)

    doc.add_paragraph()

    # ── Subject ──────────────────
    if subject:
        subj_para = doc.add_paragraph()
        subj_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subr = subj_para.add_run(f"Subject: {subject}")
        subr.font.size = Pt(14)
        subr.font.name = template.font_name
        subr.font.color.rgb = RGBColor(120, 120, 120)

    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph()

    # ── Student name ──────────────────
    if student_name:
        sname_para = doc.add_paragraph()
        sname_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        snr = sname_para.add_run(f"Submitted by: {student_name}")
        snr.font.size = Pt(13)
        snr.font.bold = True
        snr.font.name = template.font_name

    # ── Academic year ──────────────────
    from datetime import datetime
    year_para = doc.add_paragraph()
    year_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    yr = year_para.add_run(f"Academic Year: {datetime.now().year}-{datetime.now().year + 1}")
    yr.font.size = Pt(12)
    yr.font.name = template.font_name
    yr.font.color.rgb = RGBColor(100, 100, 100)

    doc.add_page_break()


def generate_college_document(
    title: str,
    topic: str,
    doc_type: str,
    template: CollegeTemplate,
    key_points: list = None,
    subject: str = None,
    student_name: str = None,
    include_references: bool = True
) -> dict:
    """
    Generate a Word document with full college branding.
    Main entry point called from route.
    """
    print(f"📄 Generating college-branded document: {title}")
    print(f"   Template: {template.college_name}")

    r, g, b = hex_to_rgb_tuple(template.primary_color)

    # Step 1 — AI plans content
    content = plan_document(
        title=title,
        topic=topic,
        doc_type=doc_type,
        key_points=key_points,
        subject=subject
    )

    # Step 2 — Create document
    doc = Document()

    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1.2)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    # Default font
    normal_style = doc.styles['Normal']
    normal_style.font.name = template.font_name
    normal_style.font.size = Pt(12)

    # Step 3 — Header and Footer
    _add_college_header(doc, template)
    _add_college_footer(doc, template)

    # Step 4 — Cover page
    _add_college_cover_page(
        doc, title, subject, student_name, doc_type, template
    )

    def add_section_heading(text, num):
        h = doc.add_heading(f"{num}. {text}", level=1)
        for run in h.runs:
            run.font.color.rgb = RGBColor(r, g, b)
            run.font.name = template.font_name
            run.font.size = Pt(14)
        h.paragraph_format.space_before = Pt(16)
        h.paragraph_format.space_after = Pt(8)

    def add_body_text(text):
        p = doc.add_paragraph(text)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(8)
        for run in p.runs:
            run.font.name = template.font_name
            run.font.size = Pt(12)

    # Step 5 — Abstract
    if content.get("abstract"):
        add_section_heading("Abstract", "")
        add_body_text(content["abstract"])
        doc.add_paragraph()

    # Step 6 — Introduction
    if content.get("introduction"):
        add_section_heading("Introduction", 1)
        add_body_text(content["introduction"])
        doc.add_paragraph()

    # Step 7 — Main sections
    for i, section_data in enumerate(content.get("sections", []), start=2):
        add_section_heading(
            section_data.get("heading", f"Section {i}"), i
        )
        add_body_text(section_data.get("content", ""))
        doc.add_paragraph()

    # Step 8 — Conclusion
    if content.get("conclusion"):
        next_num = len(content.get("sections", [])) + 2
        add_section_heading("Conclusion", next_num)
        add_body_text(content["conclusion"])
        doc.add_paragraph()

    # Step 9 — References
    if include_references and content.get("references"):
        add_section_heading("References", "")
        for i, ref in enumerate(content["references"], 1):
            ref_p = doc.add_paragraph()
            ref_p.paragraph_format.left_indent = Inches(0.3)
            ref_run = ref_p.add_run(f"[{i}]  {ref}")
            ref_run.font.size = Pt(11)
            ref_run.font.name = template.font_name

    # Step 10 — Save
    file_id = str(uuid.uuid4())
    file_name = f"{file_id}_{title[:25].replace(' ', '_')}_CollegeTemplate.docx"
    file_path = os.path.join(OUTPUT_DIR, file_name)
    doc.save(file_path)

    word_count = sum(
        len(str(v).split())
        for v in content.values()
        if isinstance(v, str)
    )

    print(f"✅ College document saved: {file_name}")

    return {
        "file_path": file_path,
        "file_name": file_name,
        "word_count": word_count,
        "template_used": template.college_name
    }